from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable, List, Optional

from docx import Document
from docx.text.paragraph import Paragraph

from .utils import (
    PARAGRAPH_DELIMITER,
    StylePalette,
    chunked,
    escape_text,
    iter_tagged_segments,
    unescape_text,
)

DOCX_SYSTEM_PROMPT = (
    "You are a meticulous translator. Translate the user content into French. "
    "Do not add or remove any tags. Each segment is wrapped in <sX>...</sX> tags. "
    "Keep tags identical and in the same order. When multiple paragraphs are provided, they are separated by "
    f"the token {PARAGRAPH_DELIMITER}. Return paragraphs in the same order using the same delimiter."
)


@dataclass(slots=True)
class ParagraphBundle:
    paragraph: Paragraph
    palette: StylePalette
    prompt: str


def iter_docx_paragraphs(document: Document) -> Iterable[Paragraph]:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_docx_paragraphs(cell)


def build_paragraph_bundle(paragraph: Paragraph) -> Optional[ParagraphBundle]:
    palette = StylePalette()
    segments: List[str] = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        style_key = palette.register_docx_run(run)
        segments.append(f"<{style_key}>{escape_text(text)}</{style_key}>")
    prompt = "".join(segments)
    if not prompt:
        return None
    return ParagraphBundle(paragraph=paragraph, palette=palette, prompt=prompt)


def clear_paragraph(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def apply_translation(bundle: ParagraphBundle, translated_payload: str) -> None:
    segments = iter_tagged_segments(translated_payload)
    if not segments:
        # If the model fails to return tagged segments, fall back to original text.
        return
    clear_paragraph(bundle.paragraph)
    for key, text in segments:
        run = bundle.paragraph.add_run(unescape_text(text))
        if bundle.palette.has(key):
            bundle.palette.apply_docx(key, run)


async def translate_docx_file(
    input_path: Path,
    output_path: Path,
    llm_client,
    progress_callback: Callable[[float, str], None] | None = None,
    batch_size: int = 8,
) -> None:
    document = Document(input_path)
    bundles = [bundle for paragraph in iter_docx_paragraphs(document) if (bundle := build_paragraph_bundle(paragraph))]

    if not bundles:
        document.save(output_path)
        return

    total = len(bundles)
    processed = 0

    for chunk in chunked(bundles, batch_size):
        prompts = [item.prompt for item in chunk]
        user_prompt = PARAGRAPH_DELIMITER.join(prompts)
        translated = await llm_client.translate(system_prompt=DOCX_SYSTEM_PROMPT, user_prompt=user_prompt)
        responses = translated.split(PARAGRAPH_DELIMITER)
        if len(responses) != len(prompts):
            # Attempt to handle mismatched responses gracefully by padding/truncating.
            responses = (responses + [responses[-1]] * len(prompts))[: len(prompts)] if responses else [translated] * len(prompts)
        for bundle, payload in zip(chunk, responses):
            apply_translation(bundle, payload)
            processed += 1
            if progress_callback:
                progress_callback(processed / total, f"Translated paragraph {processed}/{total}")

    document.save(output_path)
