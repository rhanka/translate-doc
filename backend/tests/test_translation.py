from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.storage import StorageManager
from app.translation import pipeline
from app.translation.docx import translate_docx_file


class DummyLLM:
    async def translate(self, *, system_prompt: str, user_prompt: str, temperature: float = 0.1) -> str:
        return (
            user_prompt
            .replace("Hello", "Bonjour")
            .replace("world", "monde")
            .replace("Simple text", "Texte simple")
        )


@pytest.mark.asyncio
async def test_docx_translation_preserves_runs(tmp_path: Path) -> None:
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Hello ")
    highlighted = paragraph.add_run("world")
    highlighted.bold = True
    paragraph.add_run("!")

    input_path = tmp_path / "sample.docx"
    output_path = tmp_path / "sample_translated.docx"
    doc.save(input_path)

    await translate_docx_file(input_path, output_path, DummyLLM())

    translated_doc = Document(output_path)
    translated_paragraph = translated_doc.paragraphs[0]
    assert translated_paragraph.runs[0].text == "Bonjour "
    assert translated_paragraph.runs[0].bold is None
    assert translated_paragraph.runs[1].text == "monde"
    assert translated_paragraph.runs[1].bold is True
    assert translated_paragraph.runs[2].text == "!"


@pytest.mark.asyncio
async def test_translate_job_text(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    job_id = "job1"
    filename = "note.txt"

    custom_storage = StorageManager(tmp_path)
    monkeypatch.setattr(pipeline, "storage", custom_storage)

    input_path = custom_storage.input_path(job_id, filename)
    input_path.write_text("Simple text", encoding="utf-8")

    result_path = await pipeline.translate_job(job_id, filename, DummyLLM(), progress_callback=None)

    assert result_path.exists()
    assert result_path.read_text(encoding="utf-8") == "Texte simple"
